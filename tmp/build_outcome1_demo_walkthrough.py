from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(r"D:\Codex\tmp\outcome1_demo_walkthrough")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "D4U_Outcome1_Demo_Walkthrough_GiangVien.docx"


PERSONAS = [
    (
        "SME",
        "Đăng nhu cầu thiết kế, chọn đúng Student, kiểm soát chất lượng và chỉ thanh toán khi quy trình đủ an toàn.",
        "Tạo project, xem application, gửi offer, nạp escrow qua PayOS, review Sketch/Final, rating Student.",
    ),
    (
        "Student Designer",
        "Tìm cơ hội thật, làm việc theo quy trình rõ ràng và nhận tiền minh bạch sau khi hoàn thành.",
        "Đăng ký, verify, apply, accept offer, nộp bài, nhận release vào ví, yêu cầu rút tiền, rating SME.",
    ),
    (
        "Admin / Operations",
        "Giữ hệ thống đáng tin cậy bằng kiểm duyệt, quan sát log và hỗ trợ các trường hợp vận hành nhạy cảm.",
        "Approve verification, xem trạng thái hệ thống, xử lý admin review, refund thủ công và theo dõi audit logs.",
    ),
]


DEMO_STEPS = [
    {
        "title": "1. SME tạo dự án mở để tuyển designer",
        "business_goal": "Chứng minh SME có thể biến một nhu cầu thiết kế thành một project công khai, có deadline, budget và category rõ ràng.",
        "show_ui": [
            "Trang tạo project của SME với các trường brief, category, budget, sketch deadline, final deadline, total deadline.",
            "Trạng thái ban đầu là DRAFT, sau đó publish thành OPEN.",
            "Nếu cần, nhắc nhẹ rằng AI Brief chỉ là trợ lý gợi ý, không tự publish thay SME.",
        ],
        "system_state": [
            "Project được lưu với status DRAFT rồi chuyển sang OPEN.",
            "Subscription plan được kiểm tra trước khi publish.",
            "Audit log ghi nhận việc publish project.",
        ],
        "talking_points": [
            "Điểm quan trọng ở đây là SME không chỉ đăng tin, mà đăng một workflow có cấu trúc.",
            "Hệ thống kiểm soát budget và số lượng open project theo plan để tránh lạm dụng ngay từ đầu.",
            "Từ góc nhìn sản phẩm, đây là bước chuyển từ nhu cầu mơ hồ sang cơ hội làm việc có thể match được.",
        ],
    },
    {
        "title": "2. Student nhìn thấy dự án và nộp ứng tuyển",
        "business_goal": "Cho thấy marketplace đang thực sự kết nối hai phía: Student thấy dự án phù hợp và phản hồi trực tiếp trên nền tảng.",
        "show_ui": [
            "Trang danh sách project phía Student và project detail.",
            "Nút apply hoặc quick apply theo budget/project terms.",
            "Nếu demo thêm custom proposal, cho thấy Student có thể đề xuất giá/giải pháp khác.",
        ],
        "system_state": [
            "Application được tạo với proposed price, cover letter hoặc ghi chú xác nhận theo terms.",
            "Student chưa verify sẽ không apply được.",
            "Duplicate apply cho cùng project bị chặn ở service và database.",
        ],
        "talking_points": [
            "Giá trị ở đây là mọi tương tác tuyển chọn đều nằm trong cùng một luồng sản phẩm, không phải chuyển sang chat hay email ngoài hệ thống.",
            "Rule duplicate apply giúp dữ liệu tuyển chọn sạch và công bằng hơn.",
            "Với giảng viên, đây là lúc nhấn mạnh vai trò của verification trong việc nâng chất lượng marketplace.",
        ],
    },
    {
        "title": "3. SME chọn ứng tuyển và tạo offer chính thức",
        "business_goal": "Biến một application thành cam kết làm việc có điều kiện rõ ràng trước khi thanh toán.",
        "show_ui": [
            "Danh sách application của SME.",
            "Hành động chọn application và gửi offer.",
            "Danh sách offers hoặc workspace hiển thị trạng thái WAITING_ACCEPTANCE.",
        ],
        "system_state": [
            "Offer được tạo ở trạng thái WAITING_ACCEPTANCE.",
            "Project chuyển sang OFFER_SELECTED.",
            "Offer chỉ hợp lệ trong cửa sổ 24 giờ để tránh treo vô thời hạn.",
        ],
        "talking_points": [
            "D4U tách application và offer thành hai lớp khác nhau để SME có thời gian cân nhắc trước khi thật sự đi tới cam kết.",
            "Điều này rất giống một quy trình tuyển chọn chuyên nghiệp chứ không phải chỉ là một nút chấp nhận hồ sơ.",
            "Deadline lock và expiry window là các rule quan trọng giúp quy trình có kỷ luật thời gian.",
        ],
    },
    {
        "title": "4. Student chấp nhận offer, nhưng dự án chưa bắt đầu ngay",
        "business_goal": "Cho thấy sản phẩm không coi lời hứa là đủ; project chỉ bắt đầu khi cả chấp nhận và thanh toán đều hoàn tất.",
        "show_ui": [
            "Màn hình Student offers với nút Accept/Reject.",
            "Sau khi accept, màn hình nhắc SME cần nạp escrow.",
            "Workspace phía SME hiển thị 'việc cần làm tiếp theo' là thanh toán escrow.",
        ],
        "system_state": [
            "Offer chuyển sang ACCEPTED.",
            "Project có selected student nhưng chưa IN_PROGRESS.",
            "Cửa sổ 24 giờ cho SME thanh toán được kích hoạt.",
        ],
        "talking_points": [
            "Đây là chỗ mình nói rõ trust model của D4U: accept offer chưa đủ để công việc bắt đầu.",
            "Hệ thống cố ý tách 'đồng ý hợp tác' và 'funded escrow' để bảo vệ cả hai bên.",
            "Nếu giảng viên hỏi tại sao làm phức tạp hơn một nút nhận job, đây là câu trả lời rất mạnh.",
        ],
    },
    {
        "title": "5. SME nạp escrow qua PayOS",
        "business_goal": "Chứng minh nền tảng có một điểm kiểm soát tài chính rõ ràng trước khi Student bắt đầu làm việc.",
        "show_ui": [
            "Nút thanh toán PayOS từ danh sách offer hoặc workspace.",
            "Checkout link hoặc QR từ PayOS.",
            "Sau khi webhook về, workspace refresh trạng thái dự án.",
        ],
        "system_state": [
            "Payment record được tạo trước khi gọi provider.",
            "Escrow được create hoặc reuse đúng logic.",
            "Webhook hợp lệ cập nhật payment SUCCESS, escrow FUNDED, offer ACTIVE, project IN_PROGRESS.",
        ],
        "talking_points": [
            "Ở góc business-first, mình không nói quá sâu vào code mà nhấn mạnh: tiền được khóa trước, nên Student có động lực làm thật và SME có bằng chứng thanh toán rõ ràng.",
            "Ở góc kỹ thuật, có thể bổ sung rằng client không tự đánh dấu success; backend chỉ tin webhook hợp lệ.",
            "Đây là điểm làm sản phẩm trông 'production-shaped' thay vì chỉ là demo CRUD.",
        ],
    },
    {
        "title": "6. Student nộp Sketch để SME review từng bước",
        "business_goal": "Cho thấy D4U không chờ đến cuối mới review, mà có milestone trung gian để giảm rủi ro lệch hướng.",
        "show_ui": [
            "Workspace của Student với bước nộp Sketch.",
            "Upload file submission và trạng thái SKETCH_REVIEW.",
            "Actions của SME: approve, request revision, invalid file report.",
        ],
        "system_state": [
            "Submission Sketch được tạo và project sang SKETCH_REVIEW.",
            "Notification NEW_SUBMISSION gửi cho SME.",
            "Review actions được lưu có cấu trúc; revision round được theo dõi.",
        ],
        "talking_points": [
            "Milestone Sketch khiến SME cảm thấy an tâm hơn vì không phải chờ đến cuối mới biết sản phẩm đi lệch brief.",
            "Điểm hay của Outcome 1 là review được đóng gói thành hành động nghiệp vụ rõ ràng, không phải comment trôi nổi.",
            "Nếu cần nhấn mạnh giá trị UX, có thể nói mỗi bên luôn biết 'việc tiếp theo' trong workspace.",
        ],
    },
    {
        "title": "7. Student nộp Final, hệ thống hoàn tất và giải ngân",
        "business_goal": "Khép kín vòng đời dự án: giao bài, duyệt, release tiền, ghi nhận hoàn thành.",
        "show_ui": [
            "Student submit Final sau khi Sketch đã approved.",
            "SME approve Final.",
            "Ví Student hiển thị số dư tăng lên sau release.",
        ],
        "system_state": [
            "Project sang FINAL_REVIEW rồi COMPLETED.",
            "Escrow đi qua RELEASE_PENDING tới RELEASED.",
            "Wallet Student nhận net amount sau platform fee; transaction DISBURSEMENT_CREDIT được ghi nhận.",
        ],
        "talking_points": [
            "Đây là khoảnh khắc giá trị nhất của demo: chúng ta không chỉ quản lý công việc, mà còn khép kín việc chuyển giá trị tài chính.",
            "Điểm nên nhấn mạnh là fee được đóng băng theo escrow lúc funding, tránh thay đổi bất ngờ về sau.",
            "Nếu giảng viên hỏi độ tin cậy, mình dẫn sang audit log và wallet transaction để chứng minh traceability.",
        ],
    },
    {
        "title": "8. Hai bên rating nhau và hệ thống để lại dấu vết vận hành",
        "business_goal": "Cho thấy Outcome 1 không dừng ở giao dịch đơn lẻ, mà bắt đầu hình thành niềm tin dài hạn trên nền tảng.",
        "show_ui": [
            "Màn hình rating sau project completed.",
            "Notification center của Student và SME.",
            "Nếu cần, có thể mở Swagger hoặc dữ liệu log để minh họa audit trail.",
        ],
        "system_state": [
            "Rating chỉ hợp lệ trong 7 ngày sau completion.",
            "Notifications core được tạo ở các thời điểm quan trọng.",
            "Audit logs lưu các state change quan trọng của project, payment, escrow và withdrawal.",
        ],
        "talking_points": [
            "Rating tạo ra lớp trust tiếp theo cho marketplace, còn notification giúp trải nghiệm vận hành không bị đứt quãng.",
            "Audit log là phần mình dùng để trả lời các câu hỏi khó về accountability và support sau này.",
            "Kết thúc phần này là lúc chốt rằng Outcome 1 đã hoàn thành một vòng đời marketplace có kiểm soát, chứ không chỉ là một prototype giao diện.",
        ],
    },
]


HIGHLIGHTS = [
    ("Workflow rõ ràng", "Mỗi actor đều có bước tiếp theo rõ ràng trong workspace, giảm mơ hồ khi vận hành dự án."),
    ("Role separation", "SME, Student và Admin có trách nhiệm tách biệt, giúp rule nghiệp vụ sạch và dễ kiểm soát."),
    ("Escrow trust model", "Dự án chỉ bắt đầu khi Student accept và PayOS webhook xác nhận payment thành công."),
    ("Revision / review control", "Sketch, Final, revision request và invalid file report đều là hành động có cấu trúc."),
    ("Wallet traceability", "Tiền đi từ escrow sang wallet qua disbursement có fee, net amount và transaction log rõ ràng."),
    ("Notification + audit support", "Hệ thống không chỉ chạy được mà còn để lại tín hiệu vận hành và dấu vết hỗ trợ điều tra."),
]


RULE_HIGHLIGHTS = [
    ("Subscription limit", "SME không thể publish vượt số lượng project mở theo plan hoặc vượt max budget.", "Giúp marketplace có kiểm soát và tránh lạm dụng tài nguyên."),
    ("Deadline lock", "Deadline có thể cập nhật trước khi offer được accept, nhưng bị khóa sau khi đã có cam kết.", "Bảo vệ Student khỏi việc thay đổi điều kiện sau khi đồng ý."),
    ("Duplicate apply guard", "Student chỉ apply một lần cho mỗi project.", "Dữ liệu tuyển chọn rõ ràng và dễ đánh giá hơn."),
    ("Escrow start condition", "Project chỉ start khi offer đã accepted và escrow đã funded.", "Đây là trọng tâm của trust model Outcome 1."),
    ("Wallet non-negative", "Số dư ví không được âm ở cả tầng service và database constraint.", "Bảo toàn tính đúng đắn tài chính."),
    ("Rating window", "Rating chỉ hợp lệ trong 7 ngày sau completion.", "Khuyến khích feedback kịp thời và tránh rating quá muộn."),
    ("Auditability", "Các state change quan trọng đều được ghi log.", "Giúp trả lời câu hỏi 'ai làm gì, khi nào, trạng thái đổi ra sao'."),
]


LIMITATIONS = [
    "Chưa làm realtime chat vì Outcome 1 tập trung vào core marketplace loop trước.",
    "Chưa mở full dispute workflow và dispute appeal; phần này được defer sang giai đoạn sau.",
    "Chưa có automatic bank payout; withdrawal vẫn cần Admin/Finance xử lý thủ công để giảm độ phức tạp ban đầu.",
    "Chưa mở rộng portfolio marketplace, package purchase hoàn chỉnh và AI Matching đầy đủ trong buổi demo này.",
    "Một số accepted deviation vẫn được nêu rõ trong backlog, ví dụ refund admin flow tách riêng khỏi withdrawal list.",
]


QA_SUPPORT = [
    ("Tại sao chưa có chat?", "Outcome 1 ưu tiên vòng đời giao dịch cốt lõi trước. Chat là lớp tiện ích tốt nhưng không phải thành phần bắt buộc để chứng minh marketplace vận hành an toàn."),
    ("Tại sao chọn PayOS?", "PayOS phù hợp với bối cảnh thanh toán local, có webhook rõ ràng và đủ cho mục tiêu escrow payment-in của MVP."),
    ("Làm sao bảo đảm giao dịch an toàn?", "Client không tự xác nhận payment thành công; backend chỉ tin webhook hợp lệ và mọi state change tài chính đều có audit trail."),
    ("Nếu người dùng bỏ dở giữa chừng thì sao?", "Outcome 1 đã có expiry window cho offer/payment, auto-close project trước execution, auto-detect student abandonment và admin refund flow."),
    ("Nếu có lỗi notification thì có làm hỏng giao dịch không?", "Không. Notification là non-blocking; lỗi notification không rollback business transaction chính."),
]


DEMO_CHECKLIST = [
    ("Tài khoản", "Chuẩn bị sẵn 3 profile riêng: SME, Student verified, Admin.", ""),
    ("Môi trường", "Frontend, API, database đều chạy; health check và Swagger sẵn sàng.", ""),
    ("Dữ liệu mẫu", "Có project title, brief, deadline và file submission mẫu.", ""),
    ("Payment", "PayOS return URL, cancel URL và webhook URL đã đúng môi trường demo.", ""),
    ("Link trình bày", "Có sẵn frontend URL, backend URL và route chính sẽ mở trong buổi demo.", ""),
    ("Fallback path", "Nếu live payment gặp sự cố, chuyển sang giải thích webhook discipline, workspace state và log evidence đã chuẩn bị.", ""),
    ("Q&A readiness", "Nắm rõ phần defer scope, accepted deviations và trust model để trả lời giảng viên.", ""),
]


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def ensure_style(document: Document, name: str, style_type: WD_STYLE_TYPE):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = ensure_style(document, "DemoTitle", WD_STYLE_TYPE.PARAGRAPH)
    title.base_style = normal
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(8)

    subtitle = ensure_style(document, "DemoSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = normal
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    subtitle.font.size = Pt(11)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x52, 0x5B, 0x76)
    subtitle.paragraph_format.space_after = Pt(12)

    callout = ensure_style(document, "DemoCallout", WD_STYLE_TYPE.PARAGRAPH)
    callout.base_style = normal
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    callout.font.size = Pt(10.5)
    callout.paragraph_format.space_after = Pt(4)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 14, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D4DCE5")


def set_table_layout(table, widths_in_inches):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in_inches):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(document: Document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_layout(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
        shade_cell(header_cells[index], "E8EEF5")
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_data):
            paragraph = row_cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parts = str(value).split("\n")
            for part_index, part in enumerate(parts):
                if part_index:
                    paragraph.add_run().add_break()
                paragraph.add_run(part)
    return table


def add_bullets(document: Document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_callout(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_layout(table, [6.3])
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.style = "DemoCallout"
    lead = paragraph.add_run(f"{label}: ")
    lead.bold = True
    paragraph.add_run(text)


def add_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("D4U Outcome 1 Demo Walkthrough")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph(style="DemoTitle")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run("D4U Outcome 1")
    paragraph.add_run().add_break()
    paragraph.add_run("Tài liệu demo product walkthrough với giảng viên")

    subtitle = document.add_paragraph(style="DemoSubtitle")
    subtitle.add_run(
        "Tài liệu này được viết như một script thuyết trình sản phẩm: đi từ vấn đề, "
        "giá trị cho từng vai trò, luồng demo chính, tới các điểm mạnh vận hành và "
        "những câu hỏi mà giảng viên có thể đặt ra."
    )

    add_callout(
        document,
        "Thông điệp chính",
        "D4U không chỉ là nơi đăng và nhận job thiết kế. Outcome 1 chứng minh được "
        "một vòng đời marketplace có kiểm soát: tuyển chọn, cam kết, thanh toán, "
        "review, giải ngân và ghi nhận uy tín.",
    )

    rows = [
        ("Mục tiêu buổi demo", "Cho giảng viên thấy hệ thống giải quyết bài toán thật và vận hành mạch lạc từ đầu đến cuối."),
        ("Góc trình bày", "Business-first, nhưng vẫn sẵn sàng mở sâu hơn về rule, trạng thái và trust model khi được hỏi."),
        ("Vai trò tham gia", "SME, Student Designer, Admin/Operations."),
        ("Nguồn đối chiếu", "MVP_D4U.md, README.md, D4U_CORE_INTERACTION_E2E_TEST_GUIDE_VI.md, D4U_OUTCOME1_DEPLOY_AZURE_VERCEL_VI.md, BACKLOG_D4U_MVP.md."),
    ]
    add_table(document, ["Hạng mục", "Nội dung"], rows, [1.75, 4.55])


def add_overview(document: Document) -> None:
    document.add_heading("1. D4U là gì và vì sao Outcome 1 quan trọng", level=1)
    document.add_paragraph(
        "D4U là một marketplace kết nối Student Designer với SME có nhu cầu thiết kế. "
        "Outcome 1 không cố chứng minh mọi tính năng của một sản phẩm lớn, mà tập trung "
        "vào vòng đời cốt lõi nhất: từ lúc một nhu cầu thiết kế được đăng lên, tới lúc "
        "công việc hoàn thành và tiền được giải ngân minh bạch."
    )
    add_bullets(
        document,
        [
            "Vấn đề phía SME: khó tìm đúng người, khó kiểm soát chất lượng, khó tạo niềm tin khi thanh toán.",
            "Vấn đề phía Student: thiếu cơ hội việc làm thật, thiếu quy trình rõ ràng, lo ngại bị chậm thanh toán hoặc review không minh bạch.",
            "Vai trò của D4U: biến quan hệ thuê thiết kế từ một chuỗi nhắn tin rời rạc thành một luồng nghiệp vụ có trạng thái, rule và bằng chứng vận hành.",
        ],
    )
    add_callout(
        document,
        "Điều mình muốn giảng viên nhớ",
        "Outcome 1 thành công khi người xem thấy rằng D4U đã đi qua được một vòng đời "
        "marketplace trọn vẹn, chứ không chỉ có giao diện đẹp hay các màn CRUD rời rạc.",
    )


def add_personas(document: Document) -> None:
    document.add_heading("2. Các vai trò người dùng và giá trị nhận được", level=1)
    add_table(document, ["Vai trò", "Giá trị nhận được", "Những gì sẽ thấy trong demo"], PERSONAS, [1.2, 2.5, 2.8])


def add_demo_story(document: Document) -> None:
    document.add_heading("3. Câu chuyện demo chính", level=1)
    document.add_paragraph(
        "Phần này là xương sống của buổi trình bày. Cách tốt nhất là đi theo một câu chuyện liên tục giữa SME và Student, "
        "để người nghe cảm thấy đang theo dõi một sản phẩm thật đang giải quyết một công việc thật."
    )

    for step in DEMO_STEPS:
        document.add_heading(step["title"], level=2)

        p = document.add_paragraph()
        lead = p.add_run("Mục tiêu nghiệp vụ: ")
        lead.bold = True
        p.add_run(step["business_goal"])

        document.add_heading("Những gì nên mở trên UI", level=3)
        add_bullets(document, step["show_ui"])

        document.add_heading("Trạng thái hệ thống cần nhấn mạnh", level=3)
        add_bullets(document, step["system_state"])

        add_callout(document, "Điểm để nói với giảng viên", " ".join(step["talking_points"]))


def add_highlights(document: Document) -> None:
    document.add_heading("4. Những điểm mạnh hệ thống nên highlight", level=1)
    add_table(document, ["Điểm mạnh", "Cách diễn đạt với giảng viên"], HIGHLIGHTS, [1.75, 4.55])


def add_rules(document: Document) -> None:
    document.add_heading("5. Business rules quan trọng cần kể lại bằng ngôn ngữ dễ hiểu", level=1)
    document.add_paragraph(
        "Ở phần này, mình không liệt kê khô toàn bộ rule như một test matrix. "
        "Mục tiêu là chọn đúng những rule mà khi kể ra, người nghe sẽ thấy hệ thống có kỷ luật vận hành và có chiều sâu nghiệp vụ."
    )
    add_table(document, ["Rule nổi bật", "Giải thích ngắn", "Ý nghĩa khi trình bày"], RULE_HIGHLIGHTS, [1.5, 2.35, 2.45])


def add_limitations(document: Document) -> None:
    document.add_heading("6. Giới hạn Outcome 1 và accepted deviations", level=1)
    document.add_paragraph(
        "Một buổi demo tốt không cố tỏ ra 'đã làm hết mọi thứ'. Thay vào đó, mình chủ động nói rõ những gì đã defer, "
        "để giảng viên thấy nhóm hiểu phạm vi và biết ưu tiên đúng thứ quan trọng trước."
    )
    add_bullets(document, LIMITATIONS)
    add_callout(
        document,
        "Cách chốt phần này",
        "Outcome 1 không phải phiên bản đầy đủ cuối cùng của D4U. Nó là mốc chứng minh rằng core marketplace loop đã đứng được trên chân của chính nó.",
    )


def add_qa_support(document: Document) -> None:
    document.add_heading("7. Kịch bản xử lý câu hỏi từ giảng viên", level=1)
    add_table(document, ["Câu hỏi có thể gặp", "Cách trả lời đề xuất"], QA_SUPPORT, [2.1, 4.2])


def add_checklist(document: Document) -> None:
    document.add_heading("8. Checklist trước giờ demo", level=1)
    document.add_paragraph(
        "Checklist này giúp buổi trình bày đi mượt hơn. Mỗi mục nên được kiểm tra xong trước khi bắt đầu để tránh mất nhịp trong lúc đang kể câu chuyện sản phẩm."
    )
    checklist_rows = [(title, detail, "") for title, detail, _ in DEMO_CHECKLIST]
    add_table(document, ["Hạng mục", "Việc cần sẵn sàng", "Đã kiểm tra"], checklist_rows, [1.4, 4.35, 0.75])


def add_closing(document: Document) -> None:
    document.add_heading("9. Cách kết thúc buổi trình bày", level=1)
    document.add_numbered_list = None
    add_numbered(
        document,
        [
            "Tóm lại D4U đã chứng minh được vòng đời marketplace cốt lõi: tuyển chọn, cam kết, thanh toán, review, giải ngân.",
            "Nhấn mạnh rằng Outcome 1 có trust model rõ ràng nhờ escrow qua PayOS, review theo milestone và audit trail.",
            "Chỉ ra rằng những phần chưa làm không phải lỗ hổng vô định, mà là phạm vi đã được defer có chủ đích.",
            "Kết bằng thông điệp: đây là một MVP đủ thật để bước vào giai đoạn hoàn thiện tiếp theo, không chỉ là một bản demo giao diện.",
        ],
    )


def build_document() -> Path:
    document = Document()
    set_page_layout(document)
    configure_styles(document)
    add_footer(document.sections[0])

    add_cover(document)
    add_overview(document)
    add_personas(document)
    add_demo_story(document)
    add_highlights(document)
    add_rules(document)
    add_limitations(document)
    add_qa_support(document)
    add_checklist(document)
    add_closing(document)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
