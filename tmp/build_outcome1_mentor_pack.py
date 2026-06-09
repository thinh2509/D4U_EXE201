from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(r"D:\Codex\tmp\outcome1_mentor_pack")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "D4U_Outcome1_Mentor_Review_Pack.docx"


BUSINESS_RULES = [
    ("BR-001", "Mỗi user chỉ có một role chính.", "Auth", "Bao phủ qua register/login/profile ownership."),
    ("BR-002", "Student user chỉ có một student profile.", "Profile", "Bao phủ qua create/update profile và verification flow."),
    ("BR-003", "SME user chỉ có một SME profile.", "Profile", "Bao phủ qua SME onboarding và update profile."),
    ("BR-004", "Mỗi project có tối đa một selected student.", "Offer", "Bao phủ qua chọn application, accept offer, duplicate guard."),
    ("BR-005", "Mỗi project có tối đa một escrow.", "Payment", "Bao phủ qua create/reuse escrow khi payment."),
    ("BR-006", "Student chỉ apply một lần cho mỗi project.", "Application", "Bao phủ qua duplicate application negative case."),
    ("BR-007", "SME không publish quá active open project limit.", "Subscription", "Bao phủ qua Basic/Pro limit cases."),
    ("BR-008", "SME không publish vượt max budget của plan.", "Subscription", "Bao phủ qua budget limit negative case."),
    ("BR-009", "Project chỉ start khi escrow funded và Student đã accept offer.", "Payment", "Bao phủ qua webhook success và guard start payment."),
    ("BR-010", "Student chỉ submit vào project được assign.", "Execution", "Bao phủ qua submit sketch/final ownership checks."),
    ("BR-011", "Final phải theo sau Sketch đã approved/auto-approved.", "Execution", "Bao phủ qua Final gating case."),
    ("BR-012", "Revision round phải audit được và không chặn SME yêu cầu revision.", "Review", "Bao phủ qua revision/admin review cases."),
    ("BR-013", "Final original file không được download trước release hợp lệ.", "Files", "Xác nhận trong file access và review flow."),
    ("BR-014", "Provider transaction id unique theo provider.", "Payment", "Bao phủ qua webhook/payment integrity review."),
    ("BR-015", "Wallet balance không được âm.", "Wallet", "Bao phủ qua disbursement và withdrawal negative cases."),
    ("BR-016", "Withdrawal yêu cầu wallet active và user đủ điều kiện.", "Wallet", "Bao phủ qua withdrawal precondition cases."),
    ("BR-017", "Open dispute chặn escrow release.", "Deferred", "Nêu rõ ngoài phạm vi Outcome 1, chỉ xác nhận đã defer."),
    ("BR-018", "Rating chỉ hợp lệ trong rating window.", "Rating", "Bao phủ qua case rating sau completion và 410 Gone."),
    ("BR-019", "Reputation score trong khoảng 0..100 khi được triển khai.", "Deferred", "Nêu rõ chưa thuộc Outcome 1."),
    ("BR-020", "Project status changes phải được record.", "Audit", "Bao phủ qua audit log và state transition review."),
]


FLOW_STEPS = [
    ("Bước 1", "Admin / Student / SME", "Tạo tài khoản, verify email, tạo profile và approve student verification.", "User đăng nhập được; Student ở trạng thái APPROVED; SME có hồ sơ hợp lệ."),
    ("Bước 2", "SME / Student", "SME tạo draft, publish project OPEN; Student đã verify tìm thấy project và apply.", "Project OPEN; application được lưu; duplicate apply bị chặn."),
    ("Bước 3", "SME / Student", "SME chọn application và tạo offer; Student accept trong hạn 24 giờ.", "Offer WAITING_ACCEPTANCE -> ACCEPTED; project có selected student."),
    ("Bước 4", "SME / System", "SME tạo payment PayOS; webhook hợp lệ fund escrow và start project.", "Payment SUCCESS; escrow FUNDED; offer ACTIVE; project IN_PROGRESS."),
    ("Bước 5", "Student / SME", "Student submit Sketch; SME approve hoặc request revision/invalid file.", "Project sang SKETCH_REVIEW rồi về IN_PROGRESS hoặc REVISION_REQUESTED."),
    ("Bước 6", "Student / SME / System", "Student submit Final; SME approve hoặc system auto-approve theo deadline.", "Project COMPLETED; escrow RELEASED; wallet Student nhận net amount."),
    ("Bước 7", "Student / Admin", "Smoke nhánh Student abandon trước submission hoặc withdrawal thủ công qua Admin.", "Refund / withdrawal được tạo và xử lý đúng rule."),
    ("Bước 8", "Student / SME / Reviewer", "Rating, notifications, audit logs và sign-off demo readiness.", "Rating trong 7 ngày; notification core xuất hiện; audit logs đủ bằng chứng."),
]


MODULES = [
    {
        "title": "1. Auth, Profile, Verification",
        "precondition": "Dùng môi trường local hoặc staging có API và frontend đang chạy; tài khoản Admin bootstrap sẵn.",
        "cases": [
            ("AUTH-01", "Đăng ký Student/SME, verify email, login thành công sau xác thực.", "BR-001, BR-002, BR-003; Backlog Foundation", "User PENDING trước verify; ACTIVE sau verify; /auth/me trả đúng role và profile."),
            ("AUTH-02", "Thử login trước khi xác thực email hoặc khi account bị suspended/banned.", "Backlog account status; NFR authz", "Hệ thống chặn login hoặc business action; không tạo session hợp lệ."),
            ("VER-01", "Student tạo profile, gửi verification metadata và Admin approve.", "Backlog verification", "Verification PENDING -> APPROVED; Student đủ điều kiện apply marketplace."),
        ],
    },
    {
        "title": "2. Project, Application, Offer",
        "precondition": "SME và Student đã login ở hai profile trình duyệt riêng; Student đã APPROVED.",
        "cases": [
            ("PROJ-01", "SME tạo draft hợp lệ và publish project OPEN.", "BR-007, BR-008; Backlog Project", "Project DRAFT -> OPEN; deadline và budget lưu đúng."),
            ("APP-01", "Student apply vào project OPEN; thử apply duplicate cùng project.", "BR-006; Backlog Application", "Application đầu tiên thành công; application thứ hai bị chặn."),
            ("OFFER-01", "SME tạo offer từ application và Student accept đúng hạn.", "BR-004; Offer flow", "Offer WAITING_ACCEPTANCE -> ACCEPTED; project có selected student duy nhất."),
            ("OFFER-02", "Negative: tạo offer khi sketch deadline còn dưới 48 giờ hoặc sửa deadline sau accept.", "Offer lead-time / deadline lock", "API trả conflict; deadline bị khóa sau accepted offer."),
        ],
    },
    {
        "title": "3. PayOS Payment và Expiry",
        "precondition": "Project đã có accepted offer; có cấu hình PayOS thật hoặc mock/dev theo mục đích test.",
        "cases": [
            ("PAY-01", "SME tạo payment sau khi Student accept offer.", "BR-009, BR-014; Backlog PayOS", "Payment record tạo trước; escrow create/reuse; trả checkout link hoặc QR."),
            ("PAY-02", "Webhook hợp lệ cập nhật payment SUCCESS và start project.", "BR-009, BR-014", "Payment SUCCESS; escrow FUNDED; offer ACTIVE; project IN_PROGRESS; audit log có PAYMENT_WEBHOOK_SUCCESS."),
            ("PAY-03", "Negative: webhook sai signature, amount mismatch, hoặc client tự mở success URL.", "Payment integrity", "Không đổi business state; không start project; có log phù hợp."),
            ("PAY-04", "Expiry: offer quá 24 giờ hoặc accepted offer quá 24 giờ chưa thanh toán.", "Offer/payment expiry jobs", "Offer EXPIRED; project release an toàn; pending payment hết hạn idempotent."),
        ],
    },
    {
        "title": "4. Execution, Submission, Review",
        "precondition": "Project đã IN_PROGRESS sau payment thành công.",
        "cases": [
            ("EXEC-01", "Student submit Sketch bằng file hợp lệ; SME approve Sketch.", "BR-010; Backlog Submission", "Submission SUBMITTED -> APPROVED; project SKETCH_REVIEW -> IN_PROGRESS."),
            ("EXEC-02", "SME request revision hoặc report invalid file trên submission.", "BR-012; Review actions", "Review action được lưu; revision round tăng đúng; invalid file không tăng revision round."),
            ("EXEC-03", "Student submit Final sau Sketch approved; SME approve Final.", "BR-011, BR-013", "Project FINAL_REVIEW -> COMPLETED; completed_at và rating_due_at được set."),
            ("EXEC-04", "Deadline path: total deadline auto-approve Final chờ review hoặc chuyển ADMIN_REVIEW nếu không có Final hợp lệ.", "Total deadline automation", "Kết quả đúng với trạng thái submission tại thời điểm hết hạn."),
        ],
    },
    {
        "title": "5. Student Abandon và SME Refund",
        "precondition": "Project đã IN_PROGRESS; với nhánh abandon thì chưa có submission nào.",
        "cases": [
            ("REFUND-01", "Student chủ động abandon trước submission kèm reason bắt buộc.", "Backlog Student abandon", "Project STUDENT_ABANDONED; escrow REFUND_PENDING; refund PENDING cho Admin."),
            ("REFUND-02", "Negative: Student abandon sau khi đã có submission.", "Abandon guard", "API chặn thao tác; không tạo refund mới."),
            ("REFUND-03", "Admin mark manual refund completed.", "Refund admin flow", "Refund COMPLETED; escrow REFUNDED; audit trail đầy đủ."),
        ],
    },
    {
        "title": "6. Wallet và Withdrawal",
        "precondition": "Student đã nhận disbursement hoặc có available balance phù hợp.",
        "cases": [
            ("WALLET-01", "Verify disbursement tạo wallet, tăng available balance bằng net amount.", "BR-015; Backlog Money movement", "Có DISBURSEMENT_CREDIT; wallet không âm; fee lấy từ escrow frozen rate."),
            ("WDR-01", "Student tạo payment method ngân hàng và gửi withdrawal hợp lệ.", "BR-016; Backlog Withdrawal", "Masked account number ở Student API; withdrawal PENDING; available giảm và locked tăng."),
            ("WDR-02", "Admin process withdrawal completed hoặc failed.", "Withdrawal admin processing", "WITHDRAWAL_DEBIT hoặc WITHDRAWAL_FAILED_REVERSAL được tạo; notification và audit log đúng."),
            ("WDR-03", "Negative: wallet không đủ tiền, wallet không ACTIVE, hoặc có request PENDING/PROCESSING khác.", "BR-015, BR-016", "Hệ thống chặn yêu cầu; số dư không bị lệch."),
        ],
    },
    {
        "title": "7. Rating",
        "precondition": "Project đã COMPLETED; cả SME và Student đều còn trong rating window.",
        "cases": [
            ("RATE-01", "Student rate SME và SME rate Student trong vòng 7 ngày.", "BR-018", "Mỗi bên tạo được một rating; average rating profile cập nhật."),
            ("RATE-02", "Negative: gửi rating trùng hoặc sau rating_due_at.", "BR-018", "Duplicate bị chặn; quá hạn trả 410 Gone."),
        ],
    },
    {
        "title": "8. Notifications và Audit Logs",
        "precondition": "Đã chạy ít nhất một happy path core flow từ offer tới release.",
        "cases": [
            ("NOTI-01", "Kiểm tra 5 notification core và unread count.", "Backlog Notifications", "Có NEW_OFFER, PAYMENT_SUCCESS, NEW_SUBMISSION, REVIEW_ACTION, ESCROW_RELEASED; list newest-first."),
            ("NOTI-02", "Mark read và read-all.", "Notification UX", "Unread count cập nhật đúng; không mất notification history."),
            ("AUDIT-01", "Đối chiếu audit_logs cho project/payment/escrow/wallet/withdrawal.", "BR-020; Audit backlog", "Có các action trọng yếu và machine-readable reason ở flow tự động."),
        ],
    },
    {
        "title": "9. Negative Cases Giá Trị Cao Khi Trình Mentor",
        "precondition": "Chuẩn bị dữ liệu sạch hoặc project clone để không ảnh hưởng happy path chính.",
        "cases": [
            ("NEG-01", "Basic plan publish project thứ 3 hoặc budget vượt 5,000,000 VND.", "BR-007, BR-008", "Publish bị chặn; active open project count không vượt rule."),
            ("NEG-02", "Student chưa verify apply hoặc submit vào project không được assign.", "BR-010", "Marketplace action bị chặn; không tạo submission/application trái quyền."),
            ("NEG-03", "Project quá total deadline nhưng chưa có accepted offer.", "Backlog pre-execution expiry", "Project auto-close; không còn nhận apply/offer mới; có audit và notification."),
            ("NEG-04", "Notification service lỗi nội bộ trong khi business transaction chính thành công.", "NFR reliability", "Business state vẫn commit; notification failure không rollback giao dịch."),
        ],
    },
]


SIGNOFF_ITEMS = [
    ("Scope Outcome 1", "Tất cả module in-scope đã có ít nhất 1 happy path và 1 negative case quan trọng."),
    ("Business rule coverage", "Mỗi rule trọng yếu BR-001..BR-020 đã được map hoặc được đánh dấu deferred rõ ràng."),
    ("Core 8-step flow", "Luồng SME -> Student -> PayOS -> Submission -> Release chạy xuyên suốt không blocker."),
    ("Money movement integrity", "Disbursement, withdrawal, refund không tạo double effect và không làm âm số dư."),
    ("Evidence completeness", "Mỗi test chính có status, evidence, notes, người chạy và ngày chạy."),
    ("Mentor review readiness", "Tài liệu đủ rõ để thuyết minh logic nghiệp vụ, scope, risk và trạng thái pass/fail."),
]


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def ensure_style(document: Document, name: str, style_type: WD_STYLE_TYPE):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = ensure_style(document, "MentorTitle", WD_STYLE_TYPE.PARAGRAPH)
    title.base_style = normal
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(8)

    subtitle = ensure_style(document, "MentorSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = normal
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    subtitle.font.size = Pt(10.5)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(10)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = document.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = document.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D0D7DE")


def set_table_layout(table, widths_in_inches):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in_inches):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(document: Document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_layout(table, widths)

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shade_cell(header_cells[index], "E8EEF5")
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_data):
            paragraph = row_cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = str(value)
            for part_index, part in enumerate(text.split("\n")):
                if part_index == 0:
                    paragraph.add_run(part)
                else:
                    paragraph.add_run().add_break()
                    paragraph.add_run(part)

    return table


def add_bullets(document: Document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_source_note(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "MentorSubtitle"
    paragraph.add_run(
        "Nguồn đối chiếu chính: BACKLOG_D4U_MVP.md, Requirement.md, "
        "D4U_CORE_INTERACTION_E2E_TEST_GUIDE_VI.md, D4U_COMPLETED_FEATURE_E2E_TEST_GUIDE_VI.md."
    )


def add_cover(document: Document) -> None:
    title = document.add_paragraph(style="MentorTitle")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("D4U Outcome 1 Test & Business Rule Review Pack")

    subtitle = document.add_paragraph(style="MentorSubtitle")
    subtitle.add_run(
        "Tài liệu phục vụ review với mentor, đồng thời dùng làm blank execution template "
        "cho quá trình test thủ công Outcome 1."
    )

    meta_rows = [
        ("Mục tiêu", "Trình bày phạm vi Outcome 1, business rule trọng yếu, test coverage và chỗ ghi nhận evidence."),
        ("Đối tượng đọc", "Mentor, reviewer kỹ thuật, SME/Product owner, QA lead."),
        ("Mức độ coverage", "Happy path + key negative cases; không phải exhaustive regression spec."),
        ("Trạng thái tài liệu", "Bản dùng để review và ghi nhận kết quả chạy test."),
    ]
    add_table(document, ["Hạng mục", "Nội dung"], meta_rows, [1.9, 4.6])
    add_source_note(document)


def add_summary(document: Document) -> None:
    document.add_heading("1. Tóm tắt điều hành", level=1)
    add_bullets(
        document,
        [
            "Outcome 1 tập trung vào luồng marketplace đầu tiên của D4U: từ tạo project, ứng tuyển, offer, PayOS escrow, tới submission, release và rating.",
            "Tài liệu này ưu tiên khả năng thuyết minh với mentor: mỗi nhóm test đều gắn với rule nghiệp vụ hoặc checklist backlog hiện có.",
            "Các mục ngoài Outcome 1 được đánh dấu rõ để tránh hiểu nhầm giữa feature deferred và feature lỗi/chưa hoàn tất.",
            "Mỗi test case đều để trống ô Status, Evidence, Notes, Run by và Run date để nhóm có thể dùng trực tiếp trong buổi review.",
        ],
    )


def add_scope(document: Document) -> None:
    document.add_heading("2. Outcome 1 Scope", level=1)
    document.add_heading("2.1. In scope cho buổi review", level=2)
    add_bullets(
        document,
        [
            "Foundation/Auth/Profile/Admin verification.",
            "Project marketplace, application, offer, expiry và auto-close trước execution.",
            "PayOS payment-in, webhook validation, escrow funding, expiry jobs.",
            "Execution flow: Sketch, Final, revision, invalid file, admin review, auto-approve.",
            "Student abandon, Admin manual refund, wallet, withdrawal, rating, notifications, audit logs.",
        ],
    )
    document.add_heading("2.2. Out of scope / deferred", level=2)
    add_bullets(
        document,
        [
            "Realtime chat, dispute workflow đầy đủ, dispute appeal.",
            "Portfolio Builder, Paid Packages, AI Matching, reputation score đầy đủ.",
            "Mid-project SME cancellation, partial refund by milestone, automatic bank payout, PayOS refund API.",
            "Mọi hành vi ngoài Outcome 1 nhưng đã được ghi nhận rõ trong backlog như deferred hoặc accepted deviation.",
        ],
    )


def add_business_rules(document: Document) -> None:
    document.add_heading("3. Business Rule Matrix", level=1)
    document.add_paragraph(
        "Bảng dưới đây giúp mentor nhìn nhanh từng rule nghiệp vụ chính, khu vực áp dụng và cách tài liệu test này bao phủ hoặc đánh dấu deferred."
    )
    add_table(
        document,
        ["Rule", "Mô tả ngắn", "Nhóm", "Coverage / ghi chú"],
        BUSINESS_RULES,
        [0.9, 2.5, 1.0, 2.1],
    )


def add_flow(document: Document) -> None:
    document.add_heading("4. End-to-End Demo Flow 8 Bước", level=1)
    document.add_paragraph(
        "Đây là flow chính để trình bày với mentor. Mỗi bước có actor rõ ràng và trạng thái hệ thống cần xác nhận trước khi chuyển sang bước tiếp theo."
    )
    add_table(
        document,
        ["Bước", "Actor", "Hoạt động chính", "Trạng thái kỳ vọng"],
        FLOW_STEPS,
        [0.8, 1.2, 2.2, 2.3],
    )


def add_modules(document: Document) -> None:
    document.add_heading("5. Functional Test Matrix Theo Module", level=1)
    document.add_paragraph(
        "Mỗi module dưới đây dùng một bảng checklist ngắn gọn để dễ scan khi review. Cột Thực thi được để trống cho kết quả thực tế."
    )
    for module in MODULES:
        document.add_heading(module["title"], level=2)
        precondition = document.add_paragraph()
        precondition.add_run("Precondition: ").bold = True
        precondition.add_run(module["precondition"])
        rows = []
        for case_id, objective, mapping, expected in module["cases"]:
            execution_blank = "Status:\nEvidence:\nNotes:\nRun by:\nRun date:"
            rows.append((case_id, objective, mapping, expected, execution_blank))
        add_table(
            document,
            ["TC ID", "Mục tiêu / bước tóm tắt", "Rule / mapping", "Expected result", "Thực thi"],
            rows,
            [0.65, 2.05, 1.2, 1.6, 1.0],
        )


def add_nfr_section(document: Document) -> None:
    document.add_heading("6. Các Điểm NFR Cần Nhấn Mạnh Với Mentor", level=1)
    add_bullets(
        document,
        [
            "Authorization và ownership check là lớp bảo vệ đầu tiên cho mọi business action quan trọng như publish, apply, review, withdrawal process.",
            "Payment, expiry, release và refund phải idempotent để webhook duplicate hoặc background job lặp không tạo side effect lần hai.",
            "Wallet balance không được âm cả ở tầng service validation lẫn database constraint.",
            "Notification là non-blocking; lỗi notification không được rollback business transaction chính.",
            "Audit log là nguồn bằng chứng để giải thích state transition của project, payment, escrow và withdrawal khi mentor hỏi traceability.",
        ],
    )


def add_signoff(document: Document) -> None:
    document.add_heading("7. Kết Luận Và Sign-off", level=1)
    document.add_paragraph(
        "Phần này dùng ở cuối buổi review để chốt từng khu vực đã pass, còn mở, hoặc cần follow-up thêm."
    )
    rows = [(area, criteria, "", "", "") for area, criteria in SIGNOFF_ITEMS]
    add_table(
        document,
        ["Hạng mục", "Tiêu chí pass", "Status", "Owner / người xác nhận", "Ghi chú"],
        rows,
        [1.25, 2.85, 0.65, 0.95, 0.8],
    )


def add_footer(section) -> None:
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_paragraph.add_run("D4U Outcome 1 Mentor Review Pack")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_document() -> Path:
    document = Document()
    set_page_layout(document)
    configure_styles(document)
    add_footer(document.sections[0])

    add_cover(document)
    add_summary(document)
    add_scope(document)
    add_business_rules(document)
    add_flow(document)
    add_modules(document)
    add_nfr_section(document)
    add_signoff(document)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
