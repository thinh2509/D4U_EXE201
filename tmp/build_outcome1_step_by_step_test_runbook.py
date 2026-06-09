from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(r"D:\Codex\tmp\outcome1_step_by_step_test_runbook")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "D4U_Outcome1_Test_Runbook_Step_By_Step.docx"


ACCOUNT_DATA = [
    ("Admin", "Dùng `ADMIN_EMAIL` và `ADMIN_PASSWORD` trong `.env` hoặc cấu hình deploy.", "Duyệt verification, xử lý admin review, refund, withdrawal."),
    ("SME", "Ví dụ: `sme.d4u.test+001@gmail.com` / `Sme@123456`", "Tạo project, gửi offer, thanh toán, review bài."),
    ("Student", "Ví dụ: `student.d4u.test+001@gmail.com` / `Student@123456`", "Apply, accept offer, upload file, nhận tiền, rating."),
]


PROJECT_SAMPLE = [
    ("Tiêu đề", "Outcome 1 Smoke - Brand Identity"),
    ("Brief", "Thiết kế logo, bảng màu và guideline cơ bản cho quán cà phê specialty."),
    ("Mục đích sử dụng", "Dùng cho bảng hiệu, menu, social media và bao bì."),
    ("Project type", "OPEN"),
    ("BudgetAmount", "100000"),
    ("Currency", "VND"),
    ("AllowStudentPortfolio", "true"),
    ("IsConfidential", "false"),
]


FILE_RULES = [
    ("Định dạng hợp lệ", "jpg, png, pdf"),
    ("Dung lượng tối đa", "20 MB / file"),
    ("Submission mẫu", "01 file png/jpg cho Sketch, 01 file pdf/png cho Final"),
    ("Evidence cần chụp", "Màn hình upload thành công, tên file, trạng thái submission"),
]


TEST_CONVENTIONS = [
    "Mỗi test case được thực hiện tuần tự từ trên xuống; không bỏ qua precondition.",
    "Mỗi bước đều cần đối chiếu đúng Input đã nhập và Expected Output tương ứng.",
    "Nếu có trạng thái hệ thống thay đổi, phải ghi lại status nhìn thấy trên UI; chỉ xuống Swagger/DB khi tài liệu yêu cầu.",
    "Evidence tối thiểu gồm ảnh màn hình UI, route hiện tại, badge/trạng thái, và nếu cần thì trích log hoặc response ngắn.",
    "Kết quả mỗi test case được đánh dấu một trong ba trạng thái: PASS, FAIL, BLOCKED.",
]


MODULE_CASES = [
    {
        "module": "1. Auth / Email Verification / Google Login",
        "cases": [
            {
                "id": "AUTH-01",
                "goal": "Đăng ký Student bằng email/password, xác thực email và đăng nhập thành công.",
                "precondition": "Hệ thống đang chạy; chưa tồn tại email test Student trong DB hoặc đã reset dữ liệu.",
                "evidence": "Ảnh màn hình form register, verify email, login thành công, và `/auth/me` hoặc UI current user.",
                "notes": "Đây là case mở đầu cho toàn bộ runbook. Nếu fail ở đây, các case sau với Student có thể bị BLOCKED.",
                "steps": [
                    (
                        "Mở trang đăng ký.",
                        "URL: `/register`.",
                        "Form đăng ký hiển thị đầy đủ trường email, full name, password, role.",
                    ),
                    (
                        "Nhập dữ liệu đăng ký Student.",
                        "Email: `student.d4u.test+001@gmail.com`\nFull name: `Student Outcome 1`\nPassword: `Student@123456`\nRole: `STUDENT`",
                        "Nút đăng ký khả dụng; không có validation lỗi nếu dữ liệu hợp lệ.",
                    ),
                    (
                        "Bấm nút đăng ký.",
                        "Click `Đăng ký`.",
                        "UI thông báo tài khoản đã tạo và yêu cầu xác thực email; user ở trạng thái `PENDING`.",
                    ),
                    (
                        "Thử đăng nhập trước khi verify email.",
                        "Email/password vừa đăng ký.",
                        "Hệ thống chặn đăng nhập email/password trước khi email được xác thực.",
                    ),
                    (
                        "Mở bước xác thực email.",
                        "URL: `/verify-email` hoặc flow tương đương trên UI.",
                        "Form xác thực email hiển thị email và ô nhập code.",
                    ),
                    (
                        "Nhập mã xác thực email hợp lệ.",
                        "Code nhận được qua SMTP hoặc mã theo môi trường test.",
                        "Xác thực thành công; hệ thống cho phép đăng nhập email/password.",
                    ),
                    (
                        "Đăng nhập lại sau verify.",
                        "Email/password vừa đăng ký.",
                        "Đăng nhập thành công; UI chuyển vào phiên người dùng Student.",
                    ),
                ],
            },
            {
                "id": "AUTH-02",
                "goal": "Đăng ký SME và xác minh quy tắc role công khai.",
                "precondition": "Hệ thống đang chạy; email SME chưa tồn tại.",
                "evidence": "Ảnh màn hình register SME; ảnh hoặc log cho case role invalid nếu test thêm.",
                "notes": "Có thể chạy ngay sau AUTH-01 để chuẩn bị tài khoản SME cho flow chính.",
                "steps": [
                    (
                        "Mở lại form đăng ký.",
                        "URL: `/register`.",
                        "Form đăng ký hiển thị bình thường.",
                    ),
                    (
                        "Nhập dữ liệu đăng ký SME.",
                        "Email: `sme.d4u.test+001@gmail.com`\nFull name: `SME Outcome 1`\nPassword: `Sme@123456`\nRole: `SME`",
                        "Form hợp lệ, không lỗi validation.",
                    ),
                    (
                        "Bấm đăng ký và xác thực email tương tự Student.",
                        "Mã xác thực email hợp lệ.",
                        "SME tạo thành công, xác thực thành công, đăng nhập được sau verify.",
                    ),
                    (
                        "Kiểm tra role công khai không cho `ADMIN`.",
                        "Nếu UI không có lựa chọn `ADMIN`, ghi nhận luôn; nếu test qua API/manual payload thì thử gửi role `ADMIN`.",
                        "Public registration không tạo được user role `ADMIN`.",
                    ),
                ],
            },
            {
                "id": "AUTH-03",
                "goal": "Xác minh Google login chỉ cho `STUDENT` và `SME` nếu môi trường đã bật Google OAuth.",
                "precondition": "Có `GOOGLE_AUTH_CLIENT_ID`; môi trường test cho phép Google login.",
                "evidence": "Ảnh nút Google login hoặc ghi chú `SKIPPED`/`N/A` nếu môi trường không cấu hình.",
                "notes": "Nếu môi trường local không bật Google login thì đánh dấu `N/A` thay vì FAIL.",
                "steps": [
                    (
                        "Mở trang đăng nhập.",
                        "URL: `/login`.",
                        "Nếu đã bật Google login thì có nút `Continue with Google` hoặc tương đương.",
                    ),
                    (
                        "Thực hiện đăng nhập Google với role Student hoặc SME.",
                        "Google account hợp lệ theo môi trường test.",
                        "Hệ thống chỉ tạo/cho phép session cho role `STUDENT` hoặc `SME`.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "2. Student Profile / SME Profile / Admin Verification",
        "cases": [
            {
                "id": "PROFILE-01",
                "goal": "Student tạo profile và gửi verification document metadata.",
                "precondition": "Đăng nhập bằng Student đã verify account email nhưng chưa có profile hoàn chỉnh.",
                "evidence": "Ảnh profile Student sau khi lưu; ảnh form verification; trạng thái `PENDING`.",
                "notes": "Cần tách account email verification với student verification để giải thích đúng nghiệp vụ.",
                "steps": [
                    (
                        "Mở trang profile Student.",
                        "URL: `/student/profile`.",
                        "Form profile hiển thị các trường school, major, study start year, bio.",
                    ),
                    (
                        "Nhập thông tin profile Student.",
                        "School: `ĐH Kinh tế TP.HCM`\nMajor: `Digital Design`\nStudy start year: `2022`\nBio: `Student chuyên thiết kế branding và social media.`",
                        "Nút lưu khả dụng; dữ liệu hợp lệ.",
                    ),
                    (
                        "Lưu profile Student.",
                        "Click `Lưu` hoặc `Cập nhật profile`.",
                        "Profile được lưu thành công; dữ liệu hiển thị đúng khi tải lại trang.",
                    ),
                    (
                        "Mở trang verification.",
                        "URL: `/student/verification`.",
                        "Có lựa chọn upload metadata document hoặc phương thức verification đang hỗ trợ.",
                    ),
                    (
                        "Gửi verification document metadata hợp lệ.",
                        "File metadata loại `jpg`, `png` hoặc `pdf` hợp lệ theo môi trường test.",
                        "Tạo verification request trạng thái `PENDING`.",
                    ),
                ],
            },
            {
                "id": "PROFILE-02",
                "goal": "Admin duyệt Student verification và Student trở thành `APPROVED`.",
                "precondition": "Đã có verification request `PENDING`; đăng nhập riêng bằng Admin.",
                "evidence": "Ảnh danh sách Admin verifications; chi tiết request; trạng thái Student sau approval.",
                "notes": "Nếu demo qua UI chưa đủ, có thể mở Swagger/admin endpoint đúng như guide.",
                "steps": [
                    (
                        "Đăng nhập Admin và mở danh sách verification.",
                        "URL: `/admin/verifications` hoặc dùng Swagger tương ứng.",
                        "Danh sách có request của Student vừa tạo.",
                    ),
                    (
                        "Mở chi tiết verification request.",
                        "Chọn request của `student.d4u.test+001@gmail.com`.",
                        "Admin xem được metadata document và trạng thái hiện tại là `PENDING`.",
                    ),
                    (
                        "Duyệt verification.",
                        "Thao tác approve request.",
                        "Verification chuyển sang `APPROVED`; Student profile phản ánh trạng thái mới.",
                    ),
                    (
                        "Quay lại phiên Student để kiểm tra.",
                        "Refresh `/student/profile` hoặc màn hình verification.",
                        "Student hiện đủ điều kiện marketplace, verification status là `APPROVED`.",
                    ),
                ],
            },
            {
                "id": "PROFILE-03",
                "goal": "SME tạo profile doanh nghiệp.",
                "precondition": "Đăng nhập bằng tài khoản SME đã verify email.",
                "evidence": "Ảnh màn hình profile SME sau khi lưu.",
                "notes": "Nên chạy trước khi vào module project.",
                "steps": [
                    (
                        "Mở trang profile SME.",
                        "URL: `/sme/profile`.",
                        "Form hiển thị company name, representative, phone, business field, logo tùy chọn.",
                    ),
                    (
                        "Nhập thông tin hồ sơ SME.",
                        "Company name: `Cafe Aurora`\nRepresentative: `Nguyen Minh Anh`\nPhone: `0901234567`\nBusiness field: `F&B`",
                        "Form hợp lệ, không lỗi validation.",
                    ),
                    (
                        "Lưu profile.",
                        "Click `Lưu` hoặc `Cập nhật profile`.",
                        "Profile SME được lưu thành công; khi tải lại trang vẫn còn dữ liệu.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "3. Project Create / Publish / List / Cancel",
        "cases": [
            {
                "id": "PROJECT-01",
                "goal": "SME tạo draft project với dữ liệu mẫu và publish thành công.",
                "precondition": "SME đã đăng nhập và có profile hợp lệ.",
                "evidence": "Ảnh form project; project detail ở trạng thái DRAFT và OPEN; URL chứa `projectId`.",
                "notes": "Case này là gốc cho hầu hết flow marketplace phía sau.",
                "steps": [
                    (
                        "Mở form tạo project.",
                        "URL: `/sme/projects/new`.",
                        "Trang hiển thị form tạo dự án đầy đủ.",
                    ),
                    (
                        "Nhập dữ liệu project mẫu.",
                        "Title: `Outcome 1 Smoke - Brand Identity`\nBrief: `Thiết kế logo, bảng màu và guideline cơ bản cho quán cà phê specialty.`\nUsage Purpose: `Dùng cho bảng hiệu, menu, social media và bao bì.`\nProject Type: `OPEN`\nBudget: `100000`\nCurrency: `VND`\nAllowStudentPortfolio: `true`\nIsConfidential: `false`\nSketch deadline: 1 ngày tương lai\nFinal deadline: sau sketch\nTotal deadline: sau final",
                        "Form hợp lệ; không có lỗi về budget hoặc deadline.",
                    ),
                    (
                        "Tạo draft project.",
                        "Click `Tạo draft`.",
                        "Project được tạo; UI chuyển sang detail; trạng thái hiển thị `DRAFT`.",
                    ),
                    (
                        "Publish project.",
                        "Click `Publish` trên project detail.",
                        "Project chuyển sang `OPEN`; project xuất hiện trong danh sách open projects.",
                    ),
                ],
            },
            {
                "id": "PROJECT-02",
                "goal": "Negative: kiểm tra validation deadline và budget khi tạo draft.",
                "precondition": "SME đăng nhập; đang ở form tạo project.",
                "evidence": "Ảnh validation lỗi hoặc thông báo lỗi phù hợp.",
                "notes": "Có thể chạy bằng project phụ, không dùng project happy path chính.",
                "steps": [
                    (
                        "Nhập budget không hợp lệ.",
                        "Budget: `0` hoặc số âm.",
                        "Hệ thống chặn lưu draft; hiển thị lỗi budget phải lớn hơn 0.",
                    ),
                    (
                        "Nhập deadline sai thứ tự.",
                        "Sketch deadline sau Final deadline hoặc Final deadline sau Total deadline.",
                        "Hệ thống chặn lưu draft; hiển thị lỗi về thứ tự deadline.",
                    ),
                ],
            },
            {
                "id": "PROJECT-03",
                "goal": "SME hủy project ở trạng thái được phép.",
                "precondition": "Có project ở trạng thái `DRAFT` hoặc `OPEN` chưa đi vào execution.",
                "evidence": "Ảnh trạng thái `CANCELLED` và danh sách project sau khi hủy.",
                "notes": "Nên chạy trên project phụ, tránh phá flow happy path chính.",
                "steps": [
                    (
                        "Mở project detail cần hủy.",
                        "Project trạng thái `DRAFT` hoặc `OPEN`.",
                        "Project detail hiển thị đúng status ban đầu.",
                    ),
                    (
                        "Thực hiện cancel.",
                        "Click `Hủy project` và nhập reason nếu UI yêu cầu.",
                        "Project chuyển sang `CANCELLED`; lý do hủy được lưu đúng.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "4. Application / Custom Proposal / Duplicate Guard",
        "cases": [
            {
                "id": "APP-01",
                "goal": "Student xem project OPEN và gửi quick apply thành công.",
                "precondition": "Student đã `APPROVED`; đã có project `OPEN` từ case PROJECT-01.",
                "evidence": "Ảnh project detail phía Student; thông báo apply thành công; nút apply bị disable hoặc đổi trạng thái.",
                "notes": "Dùng đúng project happy path chính.",
                "steps": [
                    (
                        "Mở danh sách project phía Student.",
                        "URL: `/student/projects`.",
                        "Project `Outcome 1 Smoke - Brand Identity` xuất hiện trong danh sách.",
                    ),
                    (
                        "Mở detail project.",
                        "Click project theo title vừa tạo.",
                        "Detail hiển thị đúng brief, budget, sketch deadline, final deadline, total deadline.",
                    ),
                    (
                        "Gửi quick apply.",
                        "Click `Gửi ứng tuyển` theo điều khoản project.",
                        "Application tạo thành công; `proposed_price` bằng budget project; UI báo đã ứng tuyển.",
                    ),
                ],
            },
            {
                "id": "APP-02",
                "goal": "Student gửi custom proposal trên project phụ.",
                "precondition": "Có một project OPEN khác dành riêng để test custom proposal.",
                "evidence": "Ảnh modal custom proposal; giá và solution note đã lưu.",
                "notes": "Không nên dùng project happy path chính để tránh thay đổi giá offer mong muốn.",
                "steps": [
                    (
                        "Mở project detail phụ.",
                        "Project OPEN thứ hai dùng cho custom proposal.",
                        "Project detail hiển thị bình thường.",
                    ),
                    (
                        "Chọn đề xuất khác.",
                        "Click `Đề xuất khác`.",
                        "Modal hoặc form custom proposal hiển thị.",
                    ),
                    (
                        "Nhập đề xuất mới.",
                        "Giá đề xuất: `120000`\nGiải pháp đề xuất: tối thiểu 20 ký tự, ví dụ `Đề xuất phương án nhận diện đậm tính social-first.`",
                        "Form hợp lệ; không lỗi validation khi đủ điều kiện.",
                    ),
                    (
                        "Gửi ứng tuyển với custom proposal.",
                        "Click `Gửi ứng tuyển`.",
                        "Application lưu đúng giá và solution note Student vừa nhập.",
                    ),
                ],
            },
            {
                "id": "APP-03",
                "goal": "Negative: chặn duplicate application cùng project.",
                "precondition": "Student đã apply vào project happy path từ APP-01.",
                "evidence": "Ảnh thông báo lỗi hoặc trạng thái UI chặn apply lần hai.",
                "notes": "Đây là business rule rất nên ghi evidence rõ.",
                "steps": [
                    (
                        "Mở lại cùng project đã apply.",
                        "Project `Outcome 1 Smoke - Brand Identity`.",
                        "UI nhận diện Student đã có application trước đó.",
                    ),
                    (
                        "Thử apply lại lần hai.",
                        "Click apply nếu nút còn khả dụng hoặc thử từ entry point liên quan.",
                        "Hệ thống chặn duplicate application; không tạo application thứ hai.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "5. Offer / Accept / Reject / Deadline Update / Expiry",
        "cases": [
            {
                "id": "OFFER-01",
                "goal": "SME chọn application và gửi offer, Student accept thành công.",
                "precondition": "Project happy path đã có application từ Student.",
                "evidence": "Ảnh danh sách applications; offer WAITING_ACCEPTANCE; offer ACCEPTED sau khi Student accept.",
                "notes": "Dùng project happy path chính.",
                "steps": [
                    (
                        "SME mở danh sách applications của project.",
                        "URL: `/sme/projects/{projectId}/applications` hoặc flow tương đương.",
                        "Danh sách hiển thị application của Student với proposed price và nội dung apply.",
                    ),
                    (
                        "SME tạo offer từ application.",
                        "Click `Chọn và gửi offer` trên application Student.",
                        "Offer được tạo ở trạng thái `WAITING_ACCEPTANCE`; project chuyển `OFFER_SELECTED`.",
                    ),
                    (
                        "Student mở danh sách offer.",
                        "URL: `/student/offers`.",
                        "Offer đúng project và đúng số tiền xuất hiện.",
                    ),
                    (
                        "Student accept offer.",
                        "Click `Chấp nhận`.",
                        "Offer chuyển sang `ACCEPTED`; project có selected student; SME thấy cần thanh toán escrow.",
                    ),
                ],
            },
            {
                "id": "OFFER-02",
                "goal": "SME cập nhật deadline trước khi Student accept và Student nhận được cập nhật.",
                "precondition": "Có offer khác đang `WAITING_ACCEPTANCE` trên project phụ.",
                "evidence": "Ảnh deadline trước/sau; notification hoặc UI update phía Student.",
                "notes": "Không chạy trên offer happy path đã ACCEPTED.",
                "steps": [
                    (
                        "SME mở project/offer phụ còn `WAITING_ACCEPTANCE`.",
                        "Project phụ với waiting offer.",
                        "Offer đang ở `WAITING_ACCEPTANCE`.",
                    ),
                    (
                        "SME cập nhật deadline.",
                        "Sketch deadline, Final deadline, Total deadline mới nhưng vẫn hợp lệ.",
                        "Deadline cập nhật thành công; waiting Student được thông báo cập nhật deadline.",
                    ),
                    (
                        "Student kiểm tra lại project/offer.",
                        "Refresh màn hình offer hoặc notification center.",
                        "Student thấy deadline mới hoặc notification `PROJECT_DEADLINES_UPDATED` tương ứng.",
                    ),
                ],
            },
            {
                "id": "OFFER-03",
                "goal": "Negative: tạo offer khi sketch deadline còn dưới 48 giờ hoặc cố sửa deadline sau accept.",
                "precondition": "SME có project phụ với deadline ngắn hoặc offer đã ACCEPTED.",
                "evidence": "Ảnh lỗi conflict hoặc chặn cập nhật deadline.",
                "notes": "Có thể chia thành hai nhánh nhỏ trong cùng case.",
                "steps": [
                    (
                        "Thử tạo offer với sketch deadline còn dưới 48 giờ.",
                        "Project phụ có sketch deadline gần hơn ngưỡng 48 giờ.",
                        "API/UI trả `409 Conflict` hoặc thông báo rằng deadline quá gần để tạo offer.",
                    ),
                    (
                        "Thử sửa deadline sau khi Student đã accept offer.",
                        "Offer/project đã ở trạng thái sau accepted.",
                        "Deadline bị khóa; hệ thống chặn cập nhật.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "6. PayOS Payment / Webhook / Payment Expiry",
        "cases": [
            {
                "id": "PAY-01",
                "goal": "SME tạo payment và webhook hợp lệ start project thành công.",
                "precondition": "Offer happy path đã `ACCEPTED`.",
                "evidence": "Ảnh nút thanh toán; checkout/QR; workspace sau khi payment thành công; trạng thái project `IN_PROGRESS`.",
                "notes": "Nếu môi trường không dùng PayOS thật, ghi rõ đang dùng mock/dev nhưng vẫn đối chiếu logic trạng thái.",
                "steps": [
                    (
                        "SME mở workspace hoặc offer list.",
                        "URL: `/sme/offers` hoặc `/projects/{projectId}/execution`.",
                        "CTA thanh toán escrow qua PayOS xuất hiện vì offer đã `ACCEPTED`.",
                    ),
                    (
                        "Tạo payment.",
                        "Click `Thanh toán PayOS`.",
                        "Payment record được tạo; escrow được create/reuse; UI nhận checkout link hoặc QR.",
                    ),
                    (
                        "Hoàn tất thanh toán hoặc mô phỏng webhook thành công.",
                        "Thực hiện theo môi trường PayOS thật hoặc mock/dev.",
                        "Payment `SUCCESS`; escrow `FUNDED`; offer `ACTIVE`; project `IN_PROGRESS`.",
                    ),
                    (
                        "Quay lại workspace Student và SME.",
                        "Refresh workspace hai bên.",
                        "Student có thể bắt đầu submit Sketch; notification `PAYMENT_SUCCESS` được tạo cho Student.",
                    ),
                ],
            },
            {
                "id": "PAY-02",
                "goal": "Negative: payment guard và client không tự set success.",
                "precondition": "Có offer phụ chưa `ACCEPTED` hoặc có payment chưa hoàn thành.",
                "evidence": "Ảnh chặn tạo payment sớm; ảnh return page không làm đổi trạng thái khi không có webhook hợp lệ.",
                "notes": "Case này thường quan trọng khi giảng giải trust model.",
                "steps": [
                    (
                        "Thử tạo payment khi offer chưa được Student accept.",
                        "Offer phụ còn `WAITING_ACCEPTANCE`.",
                        "SME không được start payment trước khi Student accept.",
                    ),
                    (
                        "Thử mở return/success page bằng tay nếu có thể.",
                        "URL success/cancel với query string client-side.",
                        "Client không tự set payment success; trạng thái backend không đổi nếu webhook chưa hợp lệ.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "7. Workspace / Sketch / Final / Revision / Invalid File / Admin Review",
        "cases": [
            {
                "id": "EXEC-01",
                "goal": "Student submit Sketch, SME approve Sketch.",
                "precondition": "Project happy path đã `IN_PROGRESS` và Student đã được chọn.",
                "evidence": "Ảnh upload file submission; trạng thái `SKETCH_REVIEW`; approve thành công.",
                "notes": "Chuẩn bị sẵn file image/pdf nhỏ hơn 20 MB.",
                "steps": [
                    (
                        "Student mở workspace project.",
                        "URL: `/projects/{projectId}/execution`.",
                        "Workspace hiển thị bước kế tiếp là submit Sketch.",
                    ),
                    (
                        "Upload file submission Sketch.",
                        "File hợp lệ `jpg/png/pdf`, <= 20 MB.",
                        "Upload thành công; file xuất hiện trong danh sách chờ submit.",
                    ),
                    (
                        "Submit Sketch.",
                        "Click `Nộp Sketch` hoặc action tương đương.",
                        "Submission tạo thành công; project chuyển `SKETCH_REVIEW`; SME nhận notification `NEW_SUBMISSION`.",
                    ),
                    (
                        "SME mở workspace và approve Sketch.",
                        "Click `Approve Sketch`.",
                        "Submission `APPROVED`; project quay về `IN_PROGRESS`; Student nhận `REVIEW_ACTION`.",
                    ),
                ],
            },
            {
                "id": "EXEC-02",
                "goal": "SME request revision hoặc report invalid file trên submission phụ.",
                "precondition": "Có submission phụ đang ở trạng thái có thể review.",
                "evidence": "Ảnh form request revision hoặc invalid file; status sau khi xử lý.",
                "notes": "Dùng project/submission phụ nếu không muốn phá flow happy path chính.",
                "steps": [
                    (
                        "SME mở submission có thể review.",
                        "Submission Sketch hoặc Final trên project phụ.",
                        "SME thấy các action review khả dụng.",
                    ),
                    (
                        "Tạo revision request.",
                        "Requested changes rõ ràng; due date hợp lệ.",
                        "Project sang `REVISION_REQUESTED`; current revision round tăng đúng.",
                    ),
                    (
                        "Hoặc report invalid file.",
                        "Reason: ví dụ `WRONG_FORMAT`; mô tả; reupload due date.",
                        "Invalid file report được lưu; không tính là revision round; Student nhận thông báo phù hợp.",
                    ),
                ],
            },
            {
                "id": "EXEC-03",
                "goal": "Student submit Final và SME approve Final để hoàn tất dự án.",
                "precondition": "Sketch đã `APPROVED` hoặc auto-approved trên project happy path.",
                "evidence": "Ảnh Final submission; project `FINAL_REVIEW`; project `COMPLETED`; wallet balance tăng.",
                "notes": "Đây là đoạn rất quan trọng của toàn bộ runbook.",
                "steps": [
                    (
                        "Student mở workspace sau khi Sketch approved.",
                        "Project happy path.",
                        "Workspace hiển thị bước kế tiếp là submit Final.",
                    ),
                    (
                        "Upload và submit Final.",
                        "File Final hợp lệ `jpg/png/pdf`, <= 20 MB.",
                        "Project chuyển `FINAL_REVIEW`; submission Final được tạo.",
                    ),
                    (
                        "SME approve Final.",
                        "Click `Approve Final`.",
                        "Submission `APPROVED`; project `COMPLETED`; `completed_at` và `rating_due_at` được set.",
                    ),
                    (
                        "Kiểm tra giải ngân.",
                        "Refresh wallet Student và transaction history.",
                        "Escrow `RELEASED`; Student wallet tăng đúng net amount; có transaction `DISBURSEMENT_CREDIT`.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "8. Student Abandon / Refund",
        "cases": [
            {
                "id": "ABANDON-01",
                "goal": "Student abandon project trước mọi submission và Admin hoàn tất refund.",
                "precondition": "Có project phụ ở trạng thái `IN_PROGRESS` nhưng chưa có submission nào.",
                "evidence": "Ảnh trạng thái `STUDENT_ABANDONED`; refund `PENDING`; refund completed bởi Admin.",
                "notes": "Không dùng project happy path chính đã đi tới completion.",
                "steps": [
                    (
                        "Student mở workspace project phụ.",
                        "Project phụ `IN_PROGRESS`, chưa có submission.",
                        "Có action abandon nếu flow hỗ trợ trên UI hoặc qua endpoint hỗ trợ test.",
                    ),
                    (
                        "Thực hiện abandon với reason bắt buộc.",
                        "Reason: `Student cannot continue this project before any submission.`",
                        "Project chuyển `STUDENT_ABANDONED`; escrow `REFUND_PENDING`; refund `PENDING` được tạo.",
                    ),
                    (
                        "Admin mở danh sách refunds.",
                        "URL/flow Admin refund tương ứng.",
                        "Refund của project phụ xuất hiện ở trạng thái `PENDING`.",
                    ),
                    (
                        "Admin mark refund completed.",
                        "Thao tác complete refund.",
                        "Refund hoàn tất; escrow chuyển `REFUNDED`.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "9. Wallet / Payment Method / Withdrawal",
        "cases": [
            {
                "id": "WALLET-01",
                "goal": "Student tạo payment method ngân hàng và gửi withdrawal hợp lệ.",
                "precondition": "Student đã có available balance đủ lớn sau disbursement hoặc seed phù hợp.",
                "evidence": "Ảnh màn hình wallet; payment method; withdrawal request; masked account number.",
                "notes": "Nếu available balance chưa đủ 50,000 VND từ happy path, dùng data phù hợp cho môi trường test hoặc note `BLOCKED`.",
                "steps": [
                    (
                        "Student mở trang ví.",
                        "URL: `/student/wallet` hoặc màn hình ví tương ứng.",
                        "Hiển thị available balance, pending/locked balance và transaction history.",
                    ),
                    (
                        "Tạo payment method ngân hàng.",
                        "Bank name: `Vietcombank`\nHolder name: `Student Outcome 1`\nAccount number: `1234567890`",
                        "Payment method lưu thành công; Student chỉ thấy masked account number khi xem lại.",
                    ),
                    (
                        "Tạo withdrawal request hợp lệ.",
                        "Amount: `50000` hoặc lớn hơn, tùy available balance.",
                        "Withdrawal tạo thành công; available giảm, locked tăng; fee 5,000 VND được tính đúng trên UI.",
                    ),
                ],
            },
            {
                "id": "WALLET-02",
                "goal": "Admin process withdrawal completed hoặc failed.",
                "precondition": "Đã có withdrawal request `PENDING` từ WALLET-01; đăng nhập Admin.",
                "evidence": "Ảnh danh sách withdrawal Admin; completed/failed state; notification phía Student.",
                "notes": "Có thể chạy completed trước, failed dùng request phụ nếu cần.",
                "steps": [
                    (
                        "Admin mở danh sách withdrawal requests.",
                        "Flow Admin withdrawal tương ứng.",
                        "Request của Student xuất hiện ở trạng thái `PENDING` hoặc `PROCESSING`.",
                    ),
                    (
                        "Process withdrawal completed.",
                        "BankTransactionReference: ví dụ `TXN-OUTCOME1-001`\nTransferredAt: thời điểm hiện tại",
                        "Withdrawal hoàn tất; tạo `WITHDRAWAL_DEBIT`; Student nhận notification completed.",
                    ),
                    (
                        "Nếu test nhánh failed trên request phụ.",
                        "Chọn thao tác failed theo flow Admin.",
                        "Tạo `WITHDRAWAL_FAILED_REVERSAL`; tiền trả về available balance.",
                    ),
                ],
            },
        ],
    },
    {
        "module": "10. Rating / Notifications / Audit Logs",
        "cases": [
            {
                "id": "RATE-01",
                "goal": "Student rating SME và SME rating Student trong rating window.",
                "precondition": "Project happy path đã `COMPLETED` và còn trong 7 ngày rating window.",
                "evidence": "Ảnh form rating hai phía; rating list hoặc profile average rating sau khi submit.",
                "notes": "Chạy cả hai chiều để khép kín Outcome 1.",
                "steps": [
                    (
                        "Student mở màn hình rating cho project completed.",
                        "Project happy path vừa completed.",
                        "Form rating khả dụng cho Student.",
                    ),
                    (
                        "Student gửi rating cho SME.",
                        "Value: `5`\nComment: `Quy trình rõ ràng, brief đầy đủ và phản hồi nhanh.`",
                        "Rating tạo thành công; profile SME cập nhật average rating phù hợp.",
                    ),
                    (
                        "SME mở form rating chiều ngược lại.",
                        "Cùng project happy path.",
                        "Form rating khả dụng cho SME.",
                    ),
                    (
                        "SME gửi rating cho Student.",
                        "Value: `5`\nComment: `Student nộp bài đúng hạn và phản hồi tốt.`",
                        "Rating tạo thành công; profile Student cập nhật average rating.",
                    ),
                ],
            },
            {
                "id": "NOTI-AUDIT-01",
                "goal": "Kiểm tra notifications core và audit logs của toàn flow.",
                "precondition": "Đã chạy ít nhất project happy path từ offer tới release.",
                "evidence": "Ảnh notification center; ảnh/unread count; query hoặc màn hình audit log nếu có.",
                "notes": "Case này nên chạy ở cuối cùng để gom bằng chứng.",
                "steps": [
                    (
                        "Mở notification center của Student.",
                        "Student đã đi qua offer, payment success, review action, escrow release.",
                        "Có các notification core: `NEW_OFFER`, `PAYMENT_SUCCESS`, `REVIEW_ACTION`, `ESCROW_RELEASED` theo ngữ cảnh.",
                    ),
                    (
                        "Mở notification center của SME.",
                        "SME của project happy path.",
                        "Có `NEW_SUBMISSION` và các notification liên quan phía SME nếu đã được tạo.",
                    ),
                    (
                        "Thực hiện mark read hoặc read-all.",
                        "Click `Đánh dấu đã đọc` hoặc `Đọc tất cả`.",
                        "Unread count giảm đúng; notification history vẫn còn.",
                    ),
                    (
                        "Đối chiếu audit logs.",
                        "Mở nguồn log phù hợp: DB query, Swagger/admin tool hoặc màn hình nếu có.",
                        "Có các log trọng yếu như `PROJECT_STATUS_CHANGED`, `PAYMENT_WEBHOOK_SUCCESS`, `ESCROW_FUNDED`, `ESCROW_RELEASED`, `WITHDRAWAL_PROCESSED`.",
                    ),
                ],
            },
        ],
    },
]


REGRESSION_SUMMARY = [
    ("Auth / Verification", "", ""),
    ("Profile / Admin approval", "", ""),
    ("Project / Publish", "", ""),
    ("Application / Offer", "", ""),
    ("Payment / Escrow", "", ""),
    ("Submission / Review", "", ""),
    ("Abandon / Refund", "", ""),
    ("Wallet / Withdrawal", "", ""),
    ("Rating", "", ""),
    ("Notifications / Audit", "", ""),
]


ACCEPTED_DEVIATIONS = [
    "Refund admin flow hiện tách riêng tại `/admin/refunds`, không gộp vào withdrawal list với `isRefund=true`.",
    "Mid-project SME cancellation, partial refund by milestone, automatic bank payout và PayOS refund API nằm ngoài Outcome 1.",
    "Portfolio Builder, AI Matching, Paid Packages, dispute workflow và realtime chat không thuộc phạm vi test chính của tài liệu này.",
]


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def ensure_style(document: Document, name: str, style_type: WD_STYLE_TYPE):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = ensure_style(document, "RunbookTitle", WD_STYLE_TYPE.PARAGRAPH)
    title.base_style = normal
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(8)

    subtitle = ensure_style(document, "RunbookSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = normal
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    subtitle.font.size = Pt(10.5)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 14, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D4DCE5")


def set_table_layout(table, widths_in_inches):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in_inches):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(document: Document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_layout(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shade_cell(header_cells[index], "E8EEF5")
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_data):
            paragraph = row_cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = str(value)
            parts = text.split("\n")
            for part_index, part in enumerate(parts):
                if part_index:
                    paragraph.add_run().add_break()
                paragraph.add_run(part)
    return table


def add_bullets(document: Document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("D4U Outcome 1 Test Runbook")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph(style="RunbookTitle")
    paragraph.add_run("D4U Outcome 1")
    paragraph.add_run().add_break()
    paragraph.add_run("Test Runbook Step-by-Step có Input/Output")

    subtitle = document.add_paragraph(style="RunbookSubtitle")
    subtitle.add_run(
        "Tài liệu này là manual UI runbook cho toàn bộ Outcome 1. "
        "Tester có thể chạy tuần tự từng bước, nhập đúng dữ liệu mẫu và đối chiếu expected output ngay trên từng dòng."
    )

    rows = [
        ("Mục tiêu", "Bao phủ toàn bộ Outcome 1 theo thứ tự thực thi thực tế, có input và output rõ trên từng bước."),
        ("Kiểu test", "Manual UI step-by-step; chỉ xuống API/DB khi thật sự cần xác minh trạng thái hoặc business rule."),
        ("Nguồn đối chiếu", "BACKLOG_D4U_MVP.md, D4U_COMPLETED_FEATURE_E2E_TEST_GUIDE_VI.md, D4U_CORE_INTERACTION_E2E_TEST_GUIDE_VI.md, README.md."),
        ("Cách dùng", "Chạy tuần tự từ phần chuẩn bị tới từng module; mỗi test case đánh dấu PASS / FAIL / BLOCKED và lưu evidence."),
    ]
    add_table(document, ["Hạng mục", "Nội dung"], rows, [1.6, 4.7])


def add_test_environment(document: Document) -> None:
    document.add_heading("1. Môi trường test và dữ liệu chuẩn bị", level=1)

    document.add_heading("1.1. Prerequisites", level=2)
    add_bullets(
        document,
        [
            ".NET 8 SDK, Docker Desktop, Git, frontend dependencies hoạt động bình thường.",
            "Có quyền chạy local backend/frontend hoặc môi trường deploy tương đương.",
            "Nếu test PayOS thật, phải có cấu hình provider, return URL, cancel URL và webhook URL hợp lệ.",
        ],
    )

    document.add_heading("1.2. Branch / Build check", level=2)
    add_numbered(
        document,
        [
            "Chạy `git status --short --branch` để xác nhận đúng branch cần test.",
            "Chạy `dotnet build` ở root để xác nhận backend build pass.",
            "Chạy `npm run lint` và `npm run build` trong `FE` để xác nhận frontend build pass.",
        ],
    )

    document.add_heading("1.3. Reset dữ liệu và health check", level=2)
    add_numbered(
        document,
        [
            "Chạy `docker compose down -v` để xóa dữ liệu cũ nếu cần clean run.",
            "Chạy `docker compose up -d --build`.",
            "Xác nhận `d4u-postgres` healthy, `d4u-api` running, `d4u-frontend` running.",
            "Mở `http://localhost:8080/health`, `http://localhost:8080/swagger`, `http://localhost:3000` và xác nhận đều truy cập được.",
        ],
    )

    document.add_heading("1.4. Tài khoản mẫu", level=2)
    add_table(document, ["Role", "Thông tin đăng nhập", "Mục đích"], ACCOUNT_DATA, [1.0, 2.8, 2.7])

    document.add_heading("1.5. Dữ liệu project mẫu", level=2)
    add_table(document, ["Trường", "Giá trị mẫu"], PROJECT_SAMPLE, [1.8, 4.7])

    document.add_heading("1.6. File mẫu upload", level=2)
    add_table(document, ["Hạng mục", "Quy ước"], FILE_RULES, [1.8, 4.7])


def add_conventions(document: Document) -> None:
    document.add_heading("2. Quy ước chạy test và ghi evidence", level=1)
    add_bullets(document, TEST_CONVENTIONS)
    add_table(
        document,
        ["Trạng thái", "Ý nghĩa", "Cách dùng"],
        [
            ("PASS", "Case chạy đúng như expected output.", "Lưu ảnh màn hình hoặc bằng chứng tối thiểu."),
            ("FAIL", "Có bước cho ra output khác expected.", "Ghi rõ bước lỗi, output thực tế và ảnh/log kèm theo."),
            ("BLOCKED", "Không thể chạy tiếp vì phụ thuộc chưa sẵn sàng.", "Ghi rõ phụ thuộc nào đang chặn và case upstream liên quan."),
        ],
        [1.0, 2.2, 3.3],
    )


def add_case(document: Document, case):
    document.add_heading(f"{case['id']} - {case['goal']}", level=3)

    meta_rows = [
        ("Mục tiêu", case["goal"]),
        ("Precondition", case["precondition"]),
        ("Evidence cần lưu", case["evidence"]),
        ("Ghi chú", case["notes"]),
    ]
    add_table(document, ["Thuộc tính", "Nội dung"], meta_rows, [1.4, 5.1])

    step_rows = []
    for index, (action, input_value, expected_output) in enumerate(case["steps"], start=1):
        step_rows.append((f"Bước {index}", action, input_value, expected_output))
    add_table(
        document,
        ["Step", "Hành động", "Input", "Expected output"],
        step_rows,
        [0.65, 1.7, 1.95, 2.2],
    )

    evidence_rows = [
        ("Status", ""),
        ("Evidence link / screenshot", ""),
        ("Output thực tế nếu khác expected", ""),
        ("Người chạy", ""),
        ("Ngày chạy", ""),
    ]
    add_table(document, ["Mục ghi nhận", "Nội dung"], evidence_rows, [1.9, 4.6])


def add_modules(document: Document) -> None:
    document.add_heading("3. Test runbook theo module", level=1)
    document.add_paragraph(
        "Phần dưới đây là khối thực thi chính. Tester nên chạy tuần tự từ module 1 tới module 10 để không bị thiếu precondition ở các flow phía sau."
    )
    for module in MODULE_CASES:
        document.add_heading(module["module"], level=2)
        for case in module["cases"]:
            add_case(document, case)


def add_regression_summary(document: Document) -> None:
    document.add_heading("4. Regression checklist và summary", level=1)
    add_table(
        document,
        ["Module", "Kết quả", "Ghi chú / blocker"],
        REGRESSION_SUMMARY,
        [2.5, 1.0, 3.0],
    )


def add_deviations(document: Document) -> None:
    document.add_heading("5. Known accepted deviations và out-of-scope", level=1)
    add_bullets(document, ACCEPTED_DEVIATIONS)


def build_document() -> Path:
    document = Document()
    set_page_layout(document)
    configure_styles(document)
    add_footer(document.sections[0])

    add_cover(document)
    add_test_environment(document)
    add_conventions(document)
    add_modules(document)
    add_regression_summary(document)
    add_deviations(document)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
